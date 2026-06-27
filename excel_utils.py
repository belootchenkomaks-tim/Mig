#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
excel_utils.py — Чтение/запись Excel-таблиц миграции.

Читает входной CSV/XLSX (столбцы A-F).
Записывает итоговую таблицу (A-R) с цветовой подсветкой.
Формирует 3 выгрузки: «Для обзвона», «Для монтажа», «Для смены вланов».
"""

import os
import csv
import re
from typing import Optional
from log_utils import Logger

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


# ── Цвета подсветки (только если openpyxl установлен) ──
if HAS_OPENPYXL:
    FILL_YELLOW = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    FILL_ORANGE = PatternFill(start_color="FFA500", end_color="FFA500", fill_type="solid")
    FILL_RED     = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
    FILL_GREEN   = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")
    HEADER_FILL = PatternFill(start_color="1E376E", end_color="1E376E", fill_type="solid")
    HEADER_FONT = Font(bold=True, color="FFFFFF", size=10)
    NORMAL_FONT = Font(size=10)
    THIN_BORDER = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin'),
    )
else:
    # Заглушки
    FILL_YELLOW = FILL_ORANGE = FILL_RED = FILL_GREEN = None
    HEADER_FILL = None
    HEADER_FONT = NORMAL_FONT = None
    THIN_BORDER = None

# ── Имена столбцов итоговой таблицы (A-R) ──
COLUMNS_A_R = [
    "Номер ОЛТ Элтекс",      # A
    "Ствол",                  # B
    "ID",                     # C
    "MAC",                    # D
    "Description",            # E
    "Уровень",                # F
    "Номер ОЛТ CData",        # G
    "Ствол, номер на стволе", # H
    "Старый влан ОЛТ",        # I
    "Адрес в US",             # J
    "Старый влан биллинг",    # K
    "Номер договора",         # L
    "Телефон",                # M
    "Новые вланы",            # N
    "S/N",                    # O
    "Формула",                # P
    "Примечание",             # Q
]

# Индексы столбцов (0-based) для подсветки
COL_J_ADDR = 9   # J — Адрес в US
COL_K_VLAN = 10  # K — Старый влан биллинг


def _format_phone(row: dict) -> str:
    """Форматировать телефон: billing (US) если отличаются (с учётом +7/8)."""
    phone = row.get("phone", "") or ""
    phone_us = row.get("_phone_us", "") or ""
    if phone_us and phone:
        def norm(p):
            # Оставляем только цифры, убираем +7/8 в начале
            digits = re.sub(r'\D', '', p)
            if len(digits) == 11 and digits[0] in ('7', '8'):
                digits = digits[1:]  # последние 10 цифр
            return digits
        if norm(phone) != norm(phone_us):
            return f"{phone} (US: {phone_us})"
    return phone or phone_us or ""


def read_input_csv(path: str, log: Logger) -> list[dict]:
    """
    Читает входной CSV/XLSX. Возвращает список словарей с ключами:
    olt, chan, id, mac, desc, rssi
    """
    ext = os.path.splitext(path)[1].lower()
    rows = []

    if ext == ".csv":
        log.info(f"Чтение CSV: {path}")
        with open(path, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f, delimiter=";")
            if not reader.fieldnames:
                log.error("CSV не содержит заголовков")
                return []
            log.debug(f"Заголовки CSV: {reader.fieldnames}")

            for i, row in enumerate(reader, 1):
                rows.append({
                    "olt": row.get("OLT", "").strip(),
                    "chan": row.get("Chan", "").strip(),
                    "id": row.get("Id", "").strip(),
                    "mac": row.get("PON Serial", "").strip(),
                    "desc": row.get("Description", "").strip(),
                    "rssi": row.get("RSSI  dbm", "").strip(),
                })
                log.debug(f"  [{i}] {rows[-1]['desc']}  MAC={rows[-1]['mac']}")

    elif ext == ".xlsx":
        if not HAS_OPENPYXL:
            log.error("Установите openpyxl: pip install openpyxl")
            return []
        log.info(f"Чтение XLSX: {path}")
        wb = openpyxl.load_workbook(path, read_only=True)
        ws = wb.active
        header = [cell.value for cell in next(ws.iter_rows(min_row=1, max_col=6))]
        for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 1):
            rows.append({
                "olt": str(row[0] or "").strip(),
                "chan": str(row[1] or "").strip(),
                "id": str(row[2] or "").strip(),
                "mac": str(row[3] or "").strip(),
                "desc": str(row[4] or "").strip(),
                "rssi": str(row[5] or "").strip(),
            })
        wb.close()
    else:
        log.error(f"Неподдерживаемый формат: {ext}")
        return []

    log.info(f"Загружено абонентов: {len(rows)}")
    return rows


_olt_ip_cache: dict = {}

def olt_name_to_ip(olt_name: str, log: Logger) -> str:
    """
    Преобразует Salsk109 → 172.18.0.109
    Salsk107 → 172.18.0.107
    Результат кэшируется — повторные вызовы не логгируются.
    """
    if olt_name in _olt_ip_cache:
        return _olt_ip_cache[olt_name]
    m = re.search(r'(\d+)$', olt_name)
    if m:
        ip = f"172.18.0.{m.group(1)}"
        log.debug(f"OLT {olt_name} → IP {ip}")
        _olt_ip_cache[olt_name] = ip
        return ip
    log.warn(f"Не удалось извлечь номер OLT из {olt_name}")
    _olt_ip_cache[olt_name] = olt_name
    return olt_name


def write_output_table(
    path: str,
    header_row: str,
    from_olt: str, from_chan: str,
    to_olt_ip: str, to_chan: str,
    data: list[dict],
    log: Logger,
) -> list[dict]:
    """
    Записывает итоговую таблицу (A-Q) с цветовой подсветкой.
    data — список словарей с ключами a..q.
    Возвращает список строк-словарей с добавленным полем 'color' для подсветки.
    """
    if not HAS_OPENPYXL:
        log.error("Установите openpyxl для записи Excel")
        return data

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Миграция"

    # ── Строка 1: заголовок "109/0 на 203/5" ──
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=17)
    cell_title = ws.cell(row=1, column=1, value=header_row)
    cell_title.font = Font(bold=True, size=12, color="1E376E")
    cell_title.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    # ── Строка 2: заголовки колонок ──
    for ci, col_name in enumerate(COLUMNS_A_R, 1):
        cell = ws.cell(row=2, column=ci, value=col_name)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = THIN_BORDER
    ws.row_dimensions[2].height = 36

    # ── Данные (начиная с 3-й строки) ──
    colors = []
    for ri, row in enumerate(data, 3):
        color = None  # по умолчанию
        # Колонка A: преобразуем Salsk109 → 172.18.0.109
        from_olt_ip = olt_name_to_ip(from_olt, log)
        vals = [
            from_olt_ip,                       # A — IP исходного OLT
            from_chan,                         # B
            row.get("id", ""),                 # C
            row.get("mac", ""),                # D
            row.get("desc", ""),               # E
            row.get("rssi", ""),               # F
            to_olt_ip,                         # G
            row.get("chan_cdata", to_chan),    # H — "5 1" (ствол + номер)
            row.get("old_vlan_olt", ""),       # I
            row.get("address_us", ""),         # J
            row.get("old_vlan_billing", ""),   # K
            row.get("contract", ""),           # L
            _format_phone(row),                # M
            row.get("new_vlan", ""),           # N
            row.get("sn", ""),                 # O
            row.get("formula", ""),            # P
            row.get("note", ""),               # Q
        ]

        # Определяем цвет строки (приоритет: красный > оранж > жёлтый)
        addr_found = row.get("address_us", "")
        vlan_billing = row.get("old_vlan_billing", "")
        vlan_olt = row.get("old_vlan_olt", "")

        if not addr_found:
            color = "yellow"
        # PPPoE (VLAN=5 или 100) — оранж не ставим, это нормально
        if not vlan_billing and vlan_olt not in ("5", "100"):
            color = "orange"
        # Сравниваем только inner VLAN (биллинг возвращает "outer:inner")
        v_bill_inner = vlan_billing.split(":")[-1] if ":" in vlan_billing else vlan_billing
        if v_bill_inner and vlan_olt and v_bill_inner != vlan_olt and vlan_olt != "?":
            color = "red"

        colors.append(color)

        fill = None
        if color == "yellow":
            fill = FILL_YELLOW
        elif color == "orange":
            fill = FILL_ORANGE
        elif color == "red":
            fill = FILL_RED
        elif color == "green":
            fill = FILL_GREEN

        for ci, val in enumerate(vals, 1):
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.font = NORMAL_FONT
            cell.border = THIN_BORDER
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            if fill:
                cell.fill = fill

    # Ширина колонок
    col_widths = [18, 8, 8, 22, 28, 12, 18, 10, 14, 30, 16, 16, 16, 14, 20, 40, 20]
    for ci, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(ci)].width = w

    wb.save(path)
    log.info(f"Итоговая таблица сохранена: {path}")
    return [dict(r, color=c) for r, c in zip(data, colors)]


def save_export_tables(base_path: str, data: list[dict], log: Logger):
    """
    Сохраняет 3 выгрузки:
    1. «Для обзвона» — D, J, M, P, Q
    2. «Для монтажа» — G, H, O, J, M, Q
    3. «Для смены вланов» — J, K, N, Q
    """
    if not HAS_OPENPYXL:
        log.error("openpyxl не установлен, пропускаем выгрузки")
        return

    def _write_export(name, columns, headers):
        path = os.path.join(os.path.dirname(base_path),
                            f"{name} {os.path.basename(base_path)}")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = name

        # Заголовки
        for ci, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=ci, value=h)
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL
            cell.border = THIN_BORDER

        # Ширина колонок для каждого типа выгрузки
        _col_widths = {
            "Для обзвона":      [22, 32, 16, 16, 25],
            "Для монтажа":      [18, 12, 22, 32, 16, 25],
            "Для смены вланов": [32, 16, 16, 14, 25],
        }
        widths = _col_widths.get(name, [18]*len(columns))
        for ci, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(ci)].width = w

        # Данные
        for ri, row in enumerate(data, 2):
            ws.row_dimensions[ri].height = 20  # чтобы текст не обрезался
            for ci, col_key in enumerate(columns, 1):
                # Телефон — с учётом billing + US
                if col_key == "phone":
                    val = _format_phone(row)
                else:
                    val = row.get(col_key, "")
                cell = ws.cell(row=ri, column=ci, value=val)
                cell.font = NORMAL_FONT
                cell.border = THIN_BORDER
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                # Переносим цвет подсветки из исходной строки
                if row.get("color") == "yellow":
                    cell.fill = FILL_YELLOW
                elif row.get("color") == "orange":
                    cell.fill = FILL_ORANGE
                elif row.get("color") == "red":
                    cell.fill = FILL_RED

        wb.save(path)
        log.info(f"Выгрузка сохранена: {path}")

    _write_export("Для обзвона",
                  ["mac", "address_us", "contract", "phone", "note"],
                  ["MAC", "Адрес в US", "Договор", "Телефон", "Примечание"])

    _write_export("Для монтажа",
                  ["olt_cdata", "chan_cdata", "sn", "address_us", "phone", "note"],
                  ["Номер ОЛТ CData", "Ствол", "S/N", "Адрес в US", "Телефон", "Примечание"])

    _write_export("Для смены вланов",
                  ["address_us", "contract", "old_vlan_billing", "new_vlan", "note"],
                  ["Адрес в US", "Договор", "Старый влан биллинг", "Новые вланы", "Примечание"])


def save_commands_to_word(path: str, commands: list[str], log: Logger):
    """Сохраняет команды в Word (.docx)."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        log.warn("python-docx не установлен, Word не создан")
        return

    doc = Document()
    doc.add_heading("Команды для C-Data OLT", level=1)

    for i, cmd in enumerate(commands, 1):
        lines = cmd.split("\n")
        for line in lines:
            p = doc.add_paragraph(line.strip())
            p.style.font.size = Pt(10)
            p.style.font.name = "Consolas"
        # Пустая строка между абонентами
        doc.add_paragraph("")

    doc.save(path)
    log.info(f"Word-файл сохранён: {path} ({len(commands)} команд)")


def save_notepad(path: str, commands: list[str], log: Logger):
    """Сохраняет текстовый файл с командами для OLT (Блокнот)."""
    with open(path, "w", encoding="utf-8") as f:
        for cmd in commands:
            f.write(cmd + "\n")
    log.info(f"Файл команд сохранён: {path} ({len(commands)} команд)")
